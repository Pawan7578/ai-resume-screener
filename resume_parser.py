# resume_parser.py
# Extract and clean text from PDF / DOCX resumes
# Handles compressed PDFs that produce spaced or merged characters

import os
import re


# ──────────────────────────────────────────────────────────────
# ARTIFACT FIXES — common compressed PDF patterns
# ──────────────────────────────────────────────────────────────
ARTIFACT_MAP = [
    # Compressed PDF spacing artifacts → correct form
    (r'M\s*y\s*S\s*Q\s*L',               'MySQL'),
    (r'M\s*o\s*n\s*g\s*o\s*D\s*B',       'MongoDB'),
    (r'J\s*A\s*V\s*A\b',                  'Java'),
    (r'P\s*H\s*P\b',                      'PHP'),
    (r'H\s*T\s*M\s*L\b',                  'HTML'),
    (r'C\s*S\s*S\b',                      'CSS'),
    (r'P\s*y\s*t\s*h\s*o\s*n',            'Python'),   # BUG FIX: removed duplicate entry below
    (r'J\s*a\s*v\s*a\s*S\s*c\s*r\s*i\s*p\s*t', 'JavaScript'),
    (r'N\s*o\s*d\s*e\s*[.\s]*j\s*s',    'Nodejs'),
    (r'G\s*i\s*t\s*H\s*u\s*b',           'GitHub'),
    (r'\bG\s*i\s*t\b',                   'Git'),
    (r'O\s*p\s*e\s*n\s*C\s*V',           'OpenCV'),
    (r'C\s*\+\s*\+',                     'C++'),
    (r'X\s*A\s*M\s*P\s*P',              'XAMPP'),
    (r'W\s*A\s*M\s*P',                   'WAMP'),
    (r'F\s*i\s*g\s*m\s*a',               'Figma'),
    (r'A\s*n\s*d\s*r\s*o\s*i\s*d',       'Android'),
    (r'L\s*i\s*n\s*u\s*x',               'Linux'),
    (r'D\s*o\s*c\s*k\s*e\s*r',           'Docker'),
    (r'A\s*W\s*S\b',                     'AWS'),
    (r'G\s*C\s*P\b',                     'GCP'),
    # CamelCase merges from over-collapsing
    (r'MachineLearning',                 'Machine Learning'),
    (r'DeepLearning',                    'Deep Learning'),
    (r'DataScience',                     'Data Science'),
    (r'DataAnalytics',                   'Data Analytics'),
    (r'JavaScript',                      'JavaScript'),
    (r'PowerBI',                         'Power BI'),
    (r'RestAPI',                         'Rest API'),
]


def fix_artifacts(text):
    """Apply regex-based artifact corrections."""
    for pattern, replacement in ARTIFACT_MAP:
        text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
    return text


def fix_spaced_chars(text):
    """
    Collapse sequences of single characters separated by spaces.
    e.g. "M a c h i n e" → "Machine"
    But preserve normal words like "in a group" (short common words).
    """
    COMMON_WORDS = {'a','i','in','is','of','to','or','at','by','an','be','do',
                    'go','he','me','my','no','on','so','up','us','we','it','if',
                    'as','am','pm','vs','pg','dj','id','ip','ui','ux','db','os'}

    lines  = text.split('\n')
    result = []

    for line in lines:
        tokens  = line.split(' ')
        i       = 0
        out     = []

        while i < len(tokens):
            # Collect a run of single/double char tokens
            run = []
            j   = i
            while j < len(tokens):
                t = tokens[j].strip()
                if len(t) <= 2 and t and t.lower() not in COMMON_WORDS:
                    run.append(tokens[j])
                    j += 1
                else:
                    break

            if len(run) >= 3:
                # This looks like spaced-out chars — collapse
                out.append(''.join(run))
                i = j
            else:
                out.append(tokens[i])
                i += 1

        result.append(' '.join(out))

    return '\n'.join(result)


def clean_extracted_text(text):
    """Full cleanup pipeline for PDF-extracted text."""
    if not text:
        return text

    # 1. Fix artifact patterns first (handles spaced-letter patterns)
    text = fix_artifacts(text)

    # 2. Collapse remaining spaced-out single characters
    text = fix_spaced_chars(text)

    # 3. Clean up whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


# ──────────────────────────────────────────────────────────────
# PDF EXTRACTION — multiple strategies with fallbacks
# ──────────────────────────────────────────────────────────────
def extract_text_from_pdf(pdf_path):
    text = ""

    # Strategy 1: pypdf
    try:
        from pypdf import PdfReader
        with open(pdf_path, 'rb') as f:
            reader = PdfReader(f)
            raw    = "\n".join(
                page.extract_text() or "" for page in reader.pages
            )
        if len(raw.strip()) > 50:
            text = raw
    except Exception as e:
        print(f"  pypdf: {e}")

    # Strategy 2: PyPDF2
    if len(text.strip()) < 50:
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                raw    = "\n".join(
                    page.extract_text() or "" for page in reader.pages
                )
            if len(raw.strip()) > len(text.strip()):
                text = raw
        except Exception as e:
            print(f"  PyPDF2: {e}")

    # Strategy 3: pdfplumber
    if len(text.strip()) < 50:
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                raw = "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
            if len(raw.strip()) > len(text.strip()):
                text = raw
        except Exception as e:
            print(f"  pdfplumber: {e}")

    # Strategy 4: pdfminer
    if len(text.strip()) < 50:
        try:
            from pdfminer.high_level import extract_text as mine
            raw = mine(pdf_path) or ""
            if len(raw.strip()) > len(text.strip()):
                text = raw
        except Exception as e:
            print(f"  pdfminer: {e}")

    if not text.strip():
        print(f"❌ Could not extract text from {os.path.basename(pdf_path)}")
        return ""

    return clean_extracted_text(text)


# ──────────────────────────────────────────────────────────────
# DOCX EXTRACTION
# ──────────────────────────────────────────────────────────────
def extract_text_from_docx(docx_path):
    text = ""
    try:
        import docx
        doc  = docx.Document(docx_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += " " + cell.text
    except Exception as e:
        print(f"⚠️  DOCX error {os.path.basename(docx_path)}: {e}")
    return text.strip()


# ──────────────────────────────────────────────────────────────
# LOAD ALL RESUMES FROM FOLDER
# ──────────────────────────────────────────────────────────────
def load_resumes_from_folder(folder_path):
    resumes    = []
    file_names = []

    for file in sorted(os.listdir(folder_path)):
        if file.startswith('.'):
            continue

        path = os.path.join(folder_path, file)
        text = ""

        if file.lower().endswith(".pdf"):
            text = extract_text_from_pdf(path)
        elif file.lower().endswith(".docx"):
            text = extract_text_from_docx(path)
        else:
            continue

        if text.strip():
            resumes.append(text)
            file_names.append(file)

            # Debug output — shows what skills were detected
            skills_found = []
            for kw in ['python','java','c++','php','html','css','javascript',
                       'node','mysql','mongodb','machine learning','opencv',
                       'git','github','cloud','r','android','figma']:
                if kw in text.lower():
                    skills_found.append(kw)
            print(f"✅ {file}: {len(text)} chars | Skills: {skills_found}")
        else:
            print(f"⚠️  Skipped (empty): {file}")

    return resumes, file_names
